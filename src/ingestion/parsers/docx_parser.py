"""
Production-grade DOCX parser using python-docx.
"""

from __future__ import annotations

import logging
import time
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public types
# ---------------------------------------------------------------------------

class HeadingLevel(int, Enum):
    H1 = 1
    H2 = 2
    H3 = 3
    H4 = 4
    H5 = 5
    H6 = 6


@dataclass
class ParseConfig:
    """Runtime configuration for the DOCX parser."""
    extract_tables: bool = True
    extract_headers_footers: bool = True
    extract_images: bool = False        # extracts image blobs + content-type
    extract_comments: bool = False
    extract_footnotes: bool = True
    extract_hyperlinks: bool = True
    preserve_heading_hierarchy: bool = True
    plain_text_only: bool = False       # collapse everything to a single string


@dataclass
class RunStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None


@dataclass
class ParsedRun:
    text: str
    style: RunStyle


@dataclass
class ParsedParagraph:
    text: str                           # plain concatenated text
    runs: list[ParsedRun]
    style_name: str                     # e.g. "Normal", "Heading 1"
    heading_level: Optional[int]        # 1–6 or None
    alignment: Optional[str]           # LEFT / CENTER / RIGHT / JUSTIFY
    list_level: Optional[int]          # 0-based indent level, None if not a list
    hyperlinks: list[str]              # URLs found in this paragraph


@dataclass
class ParsedTable:
    rows: list[list[str]]              # [row][col] plain text
    row_count: int
    col_count: int


@dataclass
class ParsedImage:
    index: int
    content_type: str                  # e.g. "image/png"
    data: bytes


@dataclass
class CoreProperties:
    title: Optional[str]
    author: Optional[str]
    last_modified_by: Optional[str]
    created: Optional[str]
    modified: Optional[str]
    description: Optional[str]
    keywords: Optional[str]
    subject: Optional[str]
    revision: Optional[int]


@dataclass
class ParseResult:
    """Complete structured output of a DOCX parse operation."""
    # Content
    paragraphs: list[ParsedParagraph]
    tables: list[ParsedTable]
    images: list[ParsedImage]
    headers: list[str]
    footers: list[str]
    footnotes: list[str]
    comments: list[str]
    # Convenience
    full_text: str                     # all paragraph text joined by newlines
    headings: list[tuple[int, str]]    # [(level, text), ...]
    # Metadata
    properties: CoreProperties
    section_count: int
    source_path: str
    processing_time_ms: float
    warnings: list[str] = field(default_factory=list)

    @property
    def success(self) -> bool:
        return bool(self.paragraphs or self.tables)

    def to_dict(self) -> dict:
        return {
            "paragraphs": [
                {
                    "text": p.text,
                    "style": p.style_name,
                    "heading_level": p.heading_level,
                    "alignment": p.alignment,
                    "list_level": p.list_level,
                    "hyperlinks": p.hyperlinks,
                    "runs": [
                        {
                            "text": r.text,
                            "bold": r.style.bold,
                            "italic": r.style.italic,
                            "underline": r.style.underline,
                        }
                        for r in p.runs
                    ],
                }
                for p in self.paragraphs
            ],
            "tables": [
                {"rows": t.rows, "row_count": t.row_count, "col_count": t.col_count}
                for t in self.tables
            ],
            "images": [
                {"index": i.index, "content_type": i.content_type}
                for i in self.images
            ],
            "headers": self.headers,
            "footers": self.footers,
            "footnotes": self.footnotes,
            "comments": self.comments,
            "full_text": self.full_text,
            "headings": [{"level": lvl, "text": txt} for lvl, txt in self.headings],
            "properties": {
                "title": self.properties.title,
                "author": self.properties.author,
                "created": self.properties.created,
                "modified": self.properties.modified,
                "description": self.properties.description,
                "keywords": self.properties.keywords,
            },
            "section_count": self.section_count,
            "source_path": self.source_path,
            "processing_time_ms": self.processing_time_ms,
            "warnings": self.warnings,
            "success": self.success,
        }


# ---------------------------------------------------------------------------
# Core parser
# ---------------------------------------------------------------------------

class DocxParser:
    """
    Production-ready DOCX parser built on python-docx.

    Features
    --------
    - Full paragraph extraction with run-level styling (bold, italic, font, size).
    - Heading hierarchy detection (H1–H6) and flat headings index.
    - List-item detection with nesting level.
    - Hyperlink URL extraction per paragraph.
    - Table extraction as 2-D string grids.
    - Header/footer text per section.
    - Footnote and endnote text.
    - Comment text via raw XML.
    - Inline image blob extraction (opt-in).
    - Core document properties (author, dates, keywords, …).
    - Structured ParseResult return type; always safe — never raises.

    Usage
    -----
    >>> parser = DocxParser()
    >>> result = parser.parse("report.docx")
    >>> print(result.full_text)
    >>> for level, heading in result.headings:
    ...     print(level, heading)

    >>> cfg = ParseConfig(extract_images=True, extract_comments=True)
    >>> parser = DocxParser(config=cfg)
    >>> result = parser.parse("contract.docx")
    >>> print(result.to_dict())
    """

    _HEADING_STYLES = {
        "heading 1": 1, "heading 2": 2, "heading 3": 3,
        "heading 4": 4, "heading 5": 5, "heading 6": 6,
    }

    def __init__(self, config: Optional[ParseConfig] = None) -> None:
        self.config = config or ParseConfig()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, file_path: str | Path) -> ParseResult:
        """
        Parse a DOCX file at *file_path*.

        Parameters
        ----------
        file_path:
            Path to a .docx file.

        Returns
        -------
        ParseResult
            Always returns a ParseResult; check `.success` for whether
            meaningful content was extracted.
        """
        path = Path(file_path)
        start = time.perf_counter()
        warnings: list[str] = []

        self._validate_path(path, warnings)
        if not path.exists():
            return self._empty_result(str(path), start, warnings)

        try:
            doc = Document(str(path))
        except Exception as exc:
            logger.error("Failed to open %s: %s", path, exc)
            warnings.append(f"Could not open file: {exc}")
            return self._empty_result(str(path), start, warnings)

        cfg = self.config

        paragraphs = self._parse_paragraphs(doc, warnings)
        tables = self._parse_tables(doc) if cfg.extract_tables else []
        headers, footers = (
            self._parse_headers_footers(doc) if cfg.extract_headers_footers
            else ([], [])
        )
        footnotes = self._parse_footnotes(doc) if cfg.extract_footnotes else []
        comments = self._parse_comments(doc) if cfg.extract_comments else []
        images = self._parse_images(doc) if cfg.extract_images else []
        properties = self._parse_properties(doc)
        section_count = len(doc.sections)

        full_text = "\n".join(p.text for p in paragraphs if p.text)
        headings = [
            (p.heading_level, p.text)
            for p in paragraphs
            if p.heading_level is not None and p.text
        ]

        elapsed = (time.perf_counter() - start) * 1000
        logger.debug(
            "Parsed %s: %d paragraphs, %d tables, %d images in %.0f ms",
            path, len(paragraphs), len(tables), len(images), elapsed,
        )

        return ParseResult(
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            headers=headers,
            footers=footers,
            footnotes=footnotes,
            comments=comments,
            full_text=full_text,
            headings=headings,
            properties=properties,
            section_count=section_count,
            source_path=str(path),
            processing_time_ms=round(elapsed, 2),
            warnings=warnings,
        )

    # ------------------------------------------------------------------
    # Paragraph extraction
    # ------------------------------------------------------------------

    def _parse_paragraphs(
        self, doc: Document, warnings: list[str]
    ) -> list[ParsedParagraph]:
        results = []
        for para in doc.paragraphs:
            parsed = self._parse_paragraph(para)
            if parsed is not None:
                results.append(parsed)
        return results

    def _parse_paragraph(self, para: Paragraph) -> Optional[ParsedParagraph]:
        style_name = para.style.name if para.style else "Normal"
        heading_level = self._heading_level(para)

        runs: list[ParsedRun] = []
        for run in para.runs:
            if not run.text:
                continue
            style = RunStyle(
                bold=bool(run.bold),
                italic=bool(run.italic),
                underline=bool(run.underline),
                strike=bool(run.font.strike),
                font_name=run.font.name,
                font_size_pt=(
                    float(run.font.size.pt) if run.font.size else None
                ),
            )
            runs.append(ParsedRun(text=run.text, style=style))

        plain = "".join(r.text for r in runs)

        # Alignment
        alignment: Optional[str] = None
        if para.alignment is not None:
            alignment = str(para.alignment).split(".")[-1]

        # List level
        list_level: Optional[int] = None
        num_pr = para._element.find(qn("w:numPr"))
        if num_pr is not None:
            ilvl = num_pr.find(qn("w:ilvl"))
            if ilvl is not None:
                list_level = int(ilvl.get(qn("w:val"), 0))

        # Hyperlinks
        hyperlinks: list[str] = []
        if self.config.extract_hyperlinks:
            hyperlinks = self._extract_hyperlinks(para)

        return ParsedParagraph(
            text=plain,
            runs=runs,
            style_name=style_name,
            heading_level=heading_level,
            alignment=alignment,
            list_level=list_level,
            hyperlinks=hyperlinks,
        )

    def _heading_level(self, para: Paragraph) -> Optional[int]:
        style_name = (para.style.name or "").lower()
        if style_name in self._HEADING_STYLES:
            return self._HEADING_STYLES[style_name]
        # Outline-level fallback via XML
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is not None:
                val = outline.get(qn("w:val"))
                if val is not None:
                    return int(val) + 1  # 0-based → 1-based
        return None

    @staticmethod
    def _extract_hyperlinks(para: Paragraph) -> list[str]:
        urls: list[str] = []
        try:
            for rel in para.part.rels.values():
                if "hyperlink" in rel.reltype:
                    urls.append(rel.target_ref)
        except Exception:
            pass
        return urls

    # ------------------------------------------------------------------
    # Table extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_tables(doc: Document) -> list[ParsedTable]:
        tables: list[ParsedTable] = []
        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            col_count = max((len(r) for r in rows), default=0)
            tables.append(ParsedTable(
                rows=rows,
                row_count=len(rows),
                col_count=col_count,
            ))
        return tables

    # ------------------------------------------------------------------
    # Headers / footers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_headers_footers(doc: Document) -> tuple[list[str], list[str]]:
        headers: list[str] = []
        footers: list[str] = []
        for section in doc.sections:
            for hdr in (section.header, section.first_page_header, section.even_page_header):
                if hdr and hdr.is_linked_to_previous is False:
                    text = "\n".join(p.text for p in hdr.paragraphs if p.text)
                    if text:
                        headers.append(text)
            for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
                if ftr and ftr.is_linked_to_previous is False:
                    text = "\n".join(p.text for p in ftr.paragraphs if p.text)
                    if text:
                        footers.append(text)
        return headers, footers

    # ------------------------------------------------------------------
    # Footnotes / endnotes
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_footnotes(doc: Document) -> list[str]:
        notes: list[str] = []
        for part_name in ("/word/footnotes.xml", "/word/endnotes.xml"):
            try:
                part = doc.part.package.part_related_by(  # type: ignore[attr-defined]
                    f"http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                    f"/{part_name.split('/')[-1].replace('.xml', '')}"
                )
                from lxml import etree
                root = part._element
                for fn in root.findall(f".//{qn('w:footnote')}") + root.findall(f".//{qn('w:endnote')}"):
                    texts = [t.text for t in fn.iter(qn("w:t")) if t.text]
                    note = "".join(texts).strip()
                    if note:
                        notes.append(note)
            except Exception:
                pass
        return notes

    # ------------------------------------------------------------------
    # Comments
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_comments(doc: Document) -> list[str]:
        comments: list[str] = []
        try:
            comments_part = doc.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            )
            root = comments_part._element
            for comment in root.findall(f".//{qn('w:comment')}"):
                texts = [t.text for t in comment.iter(qn("w:t")) if t.text]
                text = "".join(texts).strip()
                author = comment.get(qn("w:author"), "")
                if text:
                    comments.append(f"[{author}] {text}" if author else text)
        except Exception:
            pass
        return comments

    # ------------------------------------------------------------------
    # Images
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_images(doc: Document) -> list[ParsedImage]:
        images: list[ParsedImage] = []
        try:
            idx = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    part = rel.target_part
                    images.append(ParsedImage(
                        index=idx,
                        content_type=part.content_type,
                        data=part.blob,
                    ))
                    idx += 1
        except Exception as exc:
            logger.debug("Image extraction error: %s", exc)
        return images

    # ------------------------------------------------------------------
    # Core properties
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_properties(doc: Document) -> CoreProperties:
        cp = doc.core_properties
        return CoreProperties(
            title=cp.title or None,
            author=cp.author or None,
            last_modified_by=cp.last_modified_by or None,
            created=str(cp.created) if cp.created else None,
            modified=str(cp.modified) if cp.modified else None,
            description=cp.description or None,
            keywords=cp.keywords or None,
            subject=cp.subject or None,
            revision=cp.revision or None,
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _validate_path(path: Path, warnings: list[str]) -> None:
        if not path.exists():
            warnings.append(f"File not found: {path}")
        elif path.suffix.lower() not in {".docx", ".docm"}:
            warnings.append(
                f"Extension '{path.suffix}' may not be supported. Expected .docx or .docm."
            )

    @staticmethod
    def _empty_result(source: str, start: float, warnings: list[str]) -> ParseResult:
        return ParseResult(
            paragraphs=[],
            tables=[],
            images=[],
            headers=[],
            footers=[],
            footnotes=[],
            comments=[],
            full_text="",
            headings=[],
            properties=CoreProperties(
                title=None, author=None, last_modified_by=None,
                created=None, modified=None, description=None,
                keywords=None, subject=None, revision=None,
            ),
            section_count=0,
            source_path=source,
            processing_time_ms=round((time.perf_counter() - start) * 1000, 2),
            warnings=warnings,
        )